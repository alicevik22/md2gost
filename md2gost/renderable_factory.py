import logging
from functools import singledispatchmethod
from typing import Generator

from docx.shared import Parented, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from .renderable import *
from .renderable import Renderable
from . import extended_markdown


class RenderableFactory:
    def __init__(self, parent: Parented):
        self._parent = parent

    @singledispatchmethod
    def create(self, marko_element: extended_markdown.BlockElement,
               caption_info: CaptionInfo) -> Generator[Renderable, None, None]:
        paragraph = Paragraph(self._parent)
        paragraph.add_run(f"{marko_element.get_type()} не поддерживается", color=RGBColor.from_string('ff0000'))
        logging.warning(f"{marko_element.get_type()} не поддерживается")
        yield paragraph

    @staticmethod
    def _create_runs(paragraph_or_link: Paragraph | Link, children, classes: list[type] = None):
        if not classes:
            classes = []
        for child in children:
            if isinstance(child, (extended_markdown.RawText, extended_markdown.Literal)):
                paragraph_or_link.add_run(child.children,
                                          is_bold=extended_markdown.StrongEmphasis in classes or None,
                                          is_italic=extended_markdown.Emphasis in classes or None,
                                          strike_through=extended_markdown.Strikethrough in classes or None)
            elif isinstance(child, extended_markdown.CodeSpan):
                paragraph_or_link.add_run(child.children, is_italic=True)
            elif isinstance(child, (extended_markdown.LineBreak, extended_markdown.Image)):
                paragraph_or_link.add_run(" ")
            elif isinstance(child, extended_markdown.Reference):
                paragraph_or_link.add_reference(child.unique_name)
            elif isinstance(child, extended_markdown.InlineEquation):
                paragraph_or_link.add_inline_equation(child.latex_equation)
            elif isinstance(child, (extended_markdown.Link, extended_markdown.Url)):
                RenderableFactory._create_runs(paragraph_or_link.add_link_url(child.dest),
                                               child.children, classes)
            elif isinstance(child, (extended_markdown.Emphasis, extended_markdown.StrongEmphasis,
                                    extended_markdown.Strikethrough)):
                RenderableFactory._create_runs(paragraph_or_link,
                                               child.children, classes + [type(child)])
            else:
                paragraph_or_link.add_run(f" {child.get_type()} не поддерживается ",
                                          color=RGBColor.from_string("FF0000"))
                logging.warning(f"{child.get_type()} не поддерживается")

    def _create_images(self, marko_paragraph: extended_markdown.Paragraph):
        images = []
        for child in marko_paragraph.children:
            if isinstance(child, extended_markdown.Image):
                images.append(
                    Image(self._parent, child.dest,
                          CaptionInfo(child.unique_name, child.title)))
        return images

    @create.register
    def _(self, marko_paragraph: extended_markdown.Paragraph, caption_info: CaptionInfo):
        paragraph = Paragraph(self._parent)

        images = self._create_images(marko_paragraph)
        if not len(images) == len(marko_paragraph.children):
            RenderableFactory._create_runs(paragraph, marko_paragraph.children)
            yield paragraph

        yield from images

    @create.register
    def _(self, marko_heading: extended_markdown.Heading | extended_markdown.SetextHeading, caption_info: CaptionInfo):
        heading = Heading(self._parent, marko_heading.level, marko_heading.numbered)
        RenderableFactory._create_runs(heading, marko_heading.children)
        yield heading

    @create.register
    def _(self, marko_code_block: extended_markdown.FencedCode | extended_markdown.CodeBlock, caption_info: CaptionInfo):
        listing = Listing(self._parent, marko_code_block.lang, caption_info)

        text = marko_code_block.children[0].children

        listing.set_text(text)
        yield listing

    @create.register
    def _(self, marko_thematic_break: extended_markdown.ThematicBreak, caption_info: CaptionInfo):
        yield from []  # ignore

    @create.register
    def _(self, _: extended_markdown.LinkRefDef, caption_info: CaptionInfo):
        yield from []  # ignore

    @create.register
    def _(self, marko_equation: extended_markdown.Equation, caption_info: CaptionInfo):
        equation = Equation(self._parent, marko_equation.latex_equation, caption_info)
        yield equation

    @create.register
    def _(self, marko_list: extended_markdown.List, caption_info: CaptionInfo):
        list_ = List(self._parent, marko_list.ordered)

        images = []
        def create_items_from_marko(marko_list_, images, level=1):
            for list_item in marko_list_.children:
                for child in list_item.children:
                    if isinstance(child, extended_markdown.List):
                        create_items_from_marko(child, images, level + 1)
                    elif isinstance(child, extended_markdown.Paragraph):
                        item = list_.add_item(self._create_images(child), level)
                        RenderableFactory._create_runs(item, child.children)

        create_items_from_marko(marko_list, images)

        yield list_
        yield from images

    @create.register
    def _(self, marko_table: extended_markdown.Table, caption_info: CaptionInfo):
        table = Table(self._parent, len(marko_table.children), len(marko_table.children[0].children),
                      caption_info)
        for i, row in enumerate(marko_table.children):
            for j, cell in enumerate(row.children):
                paragraph = table.add_paragraph_to_cell(i, j)
                paragraph.alignment = {
                    None: None,
                    "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                    "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                    "center": WD_PARAGRAPH_ALIGNMENT.CENTER
                }[cell.align]
                RenderableFactory._create_runs(
                    paragraph,
                    cell.children
                )

        yield table

    @create.register
    def _(self, marko_toc: extended_markdown.TOC, caption_info: CaptionInfo):
        toc = ToC(self._parent)
        yield toc

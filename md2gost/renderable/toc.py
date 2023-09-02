from copy import copy
from typing import Generator

from docx.enum.text import WD_TAB_LEADER, WD_TAB_ALIGNMENT, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Parented, Pt
from docx.text.run import Run

from . import Paragraph
from .page_break import PageBreak
from .renderable import Renderable
from ..layout_tracker import LayoutState
from ..rendered_info import RenderedInfo
from ..sub_renderable import SubRenderable
from ..util import create_element


def create_field(parent: Parented, text: str, instr_text: str):
    run = Run(create_element("w:r"), None)
    run._t = create_element("w:t")

    run._element.append(create_element("w:fldChar", {
        "w:fldCharType": "begin"
    }))
    run._element.append(create_element("w:instrText", {
        "xml:space": "preserve"
    }, instr_text))
    run._element.append(create_element("w:fldChar", {
        "w:fldCharType": "separate"
    }))
    run._element.append(run._t)
    if text is not None:
        run._t.text = text
    run._element.append(create_element("w:fldChar", {
        "w:fldCharType": "end"
    }))
    return run


class ToC(Renderable):
    """
    Items are added by calling add_item(level, title, page) method.
    After the document is fully rendered fill must be called.
    """

    def __init__(self, parent: Parented):
        self._parent = parent
        self._paragraphs: list[Paragraph] = []
        self._numbering = [0 for _ in range(10)]
        pass

    def add_item(self, level: int, title: str, numbered: bool):
        """Adds items to TOC. Must be called before rendering"""
        paragraph = Paragraph(self._parent)
        paragraph._docx_paragraph.paragraph_format.tab_stops.add_tab_stop(
            paragraph._docx_paragraph.part.document.sections[0].page_width - paragraph._docx_paragraph.part.document.sections[0].left_margin
            - paragraph._docx_paragraph.part.document.sections[0].right_margin,
            alignment=WD_TAB_ALIGNMENT.RIGHT, leader=WD_TAB_LEADER.DOTS)
        paragraph._docx_paragraph.paragraph_format.tab_stops.add_tab_stop(0, alignment=WD_TAB_ALIGNMENT.LEFT,
                                                          leader=WD_TAB_LEADER.SPACES)
        paragraph._docx_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph.first_line_indent = 0

        self._numbering[level - 1] += 1
        for i in range(level, len(self._numbering)):
            self._numbering[i] = 0
        paragraph.add_run("    " * (level - 1))
        if numbered:
            paragraph.add_run(".".join([str(x) for x in self._numbering[:level]]) + ". ")
        paragraph.add_run(title)
        paragraph.add_run(f"\t?")

        self._paragraphs.append(paragraph)

    def render(self, previous_rendered: RenderedInfo, layout_state: LayoutState) \
            -> Generator[RenderedInfo | SubRenderable, None, None]:
        for paragraph in self._paragraphs:
            paragraph_rendered_infos = list(
                paragraph.render(previous_rendered, copy(layout_state)))
            layout_state.add_height(sum([info.height for info in paragraph_rendered_infos]))
            yield from paragraph_rendered_infos
            previous_rendered = paragraph_rendered_infos[-1]

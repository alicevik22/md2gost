import logging
from copy import copy
from io import BytesIO
from typing import Generator
from os import environ
import os.path

import requests
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Parented, Length
from docx.text.paragraph import Paragraph

from .caption import Caption, CaptionInfo
from .renderable import Renderable
from .requires_numbering import RequiresNumbering
from ..layout_tracker import LayoutState
from ..rendered_info import RenderedInfo
from ..util import create_element


class Image(Renderable, RequiresNumbering):
    def __init__(self, parent: Parented, path: str, caption_info: CaptionInfo | None = None):
        super().__init__("Рисунок", caption_info.unique_name if caption_info else None)
        self._parent = parent
        self._caption_info = caption_info
        self._docx_paragraph = Paragraph(create_element("w:p"), parent)
        self._docx_paragraph.paragraph_format.space_before = 0
        self._docx_paragraph.paragraph_format.space_after = 0
        self._docx_paragraph.paragraph_format.first_line_indent = 0
        self._docx_paragraph.paragraph_format.line_spacing = 1
        self._docx_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self._invalid = False

        run = self._docx_paragraph.add_run()

        if path.startswith("http"):
            bytesio = BytesIO()
            bytesio.write(requests.get(path).content)
            self._image = run.add_picture(bytesio)
        else:
            try:
                self._image = run.add_picture(path)
            except FileNotFoundError:
                logging.warning(f"Путь {path} не существует, картинка не будет добавлена")
                self._invalid = True

        self._number = None

    def set_number(self, number: int):
        self._number = number

    def resize(self, width: Length = None, height: Length = None):
        if not any((width, height)):
            return

        if not width:
            width = height * (self._image.width / self._image.height)

        if not height:
            height = width * (self._image.height / self._image.width)

        self._image.width = Length(width)
        self._image.height = Length(height)

    def render(self, previous_rendered: RenderedInfo, layout_state: LayoutState)\
            -> Generator[RenderedInfo | Renderable, None, None]:
        if self._invalid:
            yield from []
            return

        # limit width
        if self._image.width > layout_state.max_width:
            self.resize(width=layout_state.max_width)

        # limit height
        if self._image.height > layout_state.max_height:
            self.resize(height=layout_state.max_height)

        height = self._image.height

        caption = Caption(self._parent, "Рисунок", self._caption_info, self._number, False)
        caption.center()

        caption_rendered_infos = list(caption.render(None, copy(layout_state)))
        caption_height = sum([info.height for info in caption_rendered_infos])

        if height + caption_height > layout_state.remaining_page_height:
            if height * 0.7 <= (layout_state.remaining_page_height - caption_height):
                self.resize(height=layout_state.remaining_page_height - caption_height)
                height = layout_state.remaining_page_height - caption_height
            else:
                height += layout_state.remaining_page_height
                self._docx_paragraph.paragraph_format.page_break_before = True

        yield RenderedInfo(self._docx_paragraph, height)
        yield from caption_rendered_infos

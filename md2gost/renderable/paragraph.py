from typing import Generator

from docx.table import Table
from docx.oxml import CT_R
from docx.shared import Length, Parented, RGBColor, Cm
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.text.paragraph import Run as DocxRun
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml.shared import qn

from . import Renderable
from .paragraph_sizer import ParagraphSizer
from ..docx_elements import create_field
from ..layout_tracker import LayoutState
from ..util import create_element
from ..rendered_info import RenderedInfo


class Link:
    def __init__(self, docx_paragraph: DocxParagraph, style: str | None):
        self._docx_paragraph = docx_paragraph
        self._style = style

        self._hyperlink = create_element("w:hyperlink")

    def set_url(self, url: str):
        r_id = self._docx_paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        self._hyperlink.set(qn("r:id"), r_id)

    def set_anchor(self, anchor: str):
        self._hyperlink.set(qn("w:anchor"), anchor)

    def add_run(self, text: str, is_bold: bool = None, is_italic: bool = None, color: RGBColor = None,
                strike_through: bool = None):

        parts = text.split("-")
        for i, part in enumerate(parts):
            docx_run = DocxRun(create_element("w:r"), self._docx_paragraph)
            self._hyperlink.append(docx_run._element)
            docx_run.text = text
            docx_run.style = self._style
            docx_run.bold = is_bold
            docx_run.italic = is_italic
            docx_run.font.color.rgb = color
            docx_run.font.strike = strike_through
            if i != len(parts) - 1:
                self._hyperlink.append(create_element("w:r", [create_element("w:noBreakHyphen")]))

    @property
    def element(self):
        return self._hyperlink


class Reference:
    def __init__(self, unique_name: str):
        self._unique_name = unique_name
        self._element = create_field("?", f"REF {unique_name} \\h")

    @property
    def unique_name(self) -> str:
        return self._unique_name

    def set_number(self, number: int):
        self._element.xpath("w:t")[0].text = str(number)

    def element(self) -> CT_R:
        return self._element


class Paragraph(Renderable):
    def __init__(self, parent: Parented):
        self._parent = parent
        self._docx_paragraph = DocxParagraph(create_element("w:p"), parent)
        self._docx_paragraph.style = "Normal"
        self._references: list[Reference] = []

    def add_run(self, text: str, is_bold: bool = None, is_italic: bool = None, color: RGBColor = None,
                strike_through: bool = None):
        # replace all hyphens with non-breaking hyphens
        parts = text.split("-")
        for i, part in enumerate(parts):
            docx_run = self._docx_paragraph.add_run(part)
            docx_run.bold = is_bold
            docx_run.italic = is_italic
            docx_run.font.color.rgb = color
            docx_run.font.strike = strike_through
            if i != len(parts)-1:
                self._docx_paragraph.add_run()._element.\
                    append(create_element("w:noBreakHyphen"))

    @property
    def references(self) -> list[Reference]:
        return self._references

    def add_reference(self, unique_name: str):
        self._references.append(Reference(unique_name))
        self._docx_paragraph._p.append(self._references[-1].element())

    def add_link_url(self, url: str, style="Hyperlink"):
        link = Link(self._docx_paragraph, style)
        link.set_url(url)
        self._docx_paragraph._p.append(link.element)
        return link

    def add_link_anchor(self, anchor: str, style="Hyperlink"):
        link = Link(self._docx_paragraph, style)
        link.set_anchor(anchor)
        self._docx_paragraph._p.append(link.element)
        return link

    def add_inline_equation(self, formula: str):
        # omml = inline_omml(latex_to_omml(formula))
        # for r in omml.xpath("//m:r", namespaces=omml.nsmap):
        #     r.append(create_element("w:rPr", [
        #         create_element("w:sz", {"w:val": "24"}),
        #         create_element("w:szCs", {"w:val": "24"}),
        #     ]))
        # self._docx_paragraph._element.append(omml)
        self.add_run(formula, is_italic=True)

    @property
    def page_break_before(self) -> bool:
        return self._docx_paragraph.paragraph_format.page_break_before

    @page_break_before.setter
    def page_break_before(self, value: bool):
        self._docx_paragraph.paragraph_format.page_break_before = value

    @property
    def style(self):
        return self._docx_paragraph.style

    @style.setter
    def style(self, value: str):
        self._docx_paragraph.style = value

    @property
    def first_line_indent(self):
        return self._docx_paragraph.paragraph_format.first_line_indent

    @first_line_indent.setter
    def first_line_indent(self, value: Length):
        self._docx_paragraph.paragraph_format.first_line_indent = value

    @property
    def alignment(self):
        return self._docx_paragraph.alignment

    @alignment.setter
    def alignment(self, value: WD_PARAGRAPH_ALIGNMENT):
        self._docx_paragraph.alignment = value

    def render(self, previous_rendered: RenderedInfo, layout_state: LayoutState)\
            -> Generator[RenderedInfo | Renderable, None, None]:
        remaining_space = layout_state.remaining_page_height

        # add space before if the previous element is table
        if isinstance(previous_rendered, RenderedInfo) and isinstance(previous_rendered.docx_element, Table):
            self._docx_paragraph.paragraph_format.space_before = Cm(0.35)  # todo: remake

        if self.page_break_before:
            layout_state.add_height(layout_state.remaining_page_height)

        height_data = ParagraphSizer(
            self._docx_paragraph,
            previous_rendered.docx_element
                      if previous_rendered and isinstance(previous_rendered.docx_element, DocxParagraph) else None,
                      layout_state.max_width).calculate_height()

        if layout_state.current_page_height == 0 and layout_state.page > 1:
            height_data.before = 0

        fitting_lines = 0
        for lines in range(1, height_data.lines+1):
            if height_data.before + ((lines - 1) * height_data.line_spacing + 1) * height_data.line_height \
                    > layout_state.remaining_page_height:
                break
            fitting_lines += 1

        if fitting_lines == height_data.lines:
            # the whole paragraph fits page
            height = min(height_data.full, layout_state.remaining_page_height)
        elif fitting_lines <= 1 or (height_data.lines-fitting_lines == 1 and height_data.lines == 3):
            # if only no or only one line fits the page, paragraph goes to the next page
            height = layout_state.remaining_page_height + height_data.full
        elif height_data.lines-fitting_lines == 1:
            # if all lines except last fit the page, the last two lines go to the new page
            height = layout_state.remaining_page_height + \
                     height_data.before + height_data.line_height * height_data.line_spacing * 2 \
                     + height_data.after
        else:
            height = layout_state.remaining_page_height + \
                     height_data.before + height_data.line_height * height_data.line_spacing * \
                     (height_data.lines-fitting_lines) + height_data.after

        if self.page_break_before:
            height += remaining_space

        yield (previous_rendered := RenderedInfo(self._docx_paragraph, Length(height)))
        layout_state.add_height(height)
